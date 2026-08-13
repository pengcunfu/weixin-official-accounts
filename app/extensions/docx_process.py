import uuid
from typing import Dict, List, Any, Tuple, Optional
from datetime import datetime
from PIL import Image
from io import BytesIO
import zipfile
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
import re
from docx import Document


class DocxProcess:
    """DOCX文档处理工具类"""

    def __init__(self):
        """初始化，获取配置并确保目录存在"""
        from app.extensions.config import config_manager

        # 获取配置
        self.images_dir = config_manager.get('upload.folders.images')
        self.cover_dir = config_manager.get('upload.folders.cover')

        # 确保目录存在
        os.makedirs(self.images_dir, exist_ok=True)
        os.makedirs(self.cover_dir, exist_ok=True)

        # 支持的图片格式
        self.image_extensions = {'.png', '.jpg',
                                 '.jpeg', '.gif', '.bmp', '.tiff'}

    def _save_cover_image(self, image_info: Dict) -> str:
        """将图片保存为封面到指定目录"""
        try:
            import shutil

            # 生成封面文件名
            original_filename = image_info.get('filename', '')
            ext = os.path.splitext(original_filename)[1]
            cover_filename = f"cover_{uuid.uuid4()}{ext}"

            # 源文件路径（images目录中的文件）
            source_path = image_info.get('absolute_path')
            # 目标路径（cover目录）
            target_path = os.path.join(self.cover_dir, cover_filename)

            if source_path and os.path.exists(source_path):
                # 复制图片到封面目录
                shutil.copy2(source_path, target_path)

                # 直接使用cover_dir拼接
                return f"{self.cover_dir}/{cover_filename}"

            return None

        except Exception as e:
            print(f"保存封面图片失败: {str(e)}")
            return None

    def parse_docx(self, file_path: str, current_user=None, relative_path: str = "", file_ext: str = "docx"):
        """
        解析docx文件，直接创建并返回Article对象

        Args:
            file_path: docx文件路径
            current_user: 当前用户对象
            relative_path: 文件的相对路径
            file_ext: 文件扩展名

        Returns:
            Article: 创建的Article对象
        """
        from app.models.article import Article
        import json

        try:
            print(f"开始解析文档: {file_path}")

            # 打开文档
            doc = Document(file_path)

            # 提取基本信息
            metadata = self._extract_metadata(doc, file_path)

            # 提取内容和图片
            content_html, images_info = self._extract_content_and_images(
                doc, file_path)

            # 处理封面图片
            cover_path = None
            if images_info:
                # 使用第一张图片作为封面
                first_image = images_info[0]
                cover_path = self._save_cover_image(first_image)

            # 计算字数
            word_count = self._calculate_word_count(content_html)

            # 创建Article对象
            new_article = Article(
                title=metadata.get('title') or os.path.splitext(
                    os.path.basename(file_path))[0],
                category=metadata.get('category') or '默认',
                file_type=file_ext,
                file_path=relative_path,
                cover_path=cover_path,  # 使用第一张图片作为封面
                status='草稿',
                saved_status='未存稿',
                public_account_nickname=None,
                author_nickname=metadata.get('author') or None,
                likes=0,
                uploader_phone=current_user.phone if current_user else None,
                word_count=word_count,
                author_id=current_user.id if current_user else None,
                draft_media_id=None,
                saved_time=None,
                content_html=content_html
            )
            
            # 使用模型方法设置图片信息
            new_article.set_images_info(images_info)

            print(f"文档解析完成，创建文章: {new_article.title}")
            return new_article

        except Exception as e:
            print(f"文档解析失败: {str(e)}")
            # 返回一个基本的Article对象，即使解析失败
            from app.models.article import Article
            return Article(
                title=os.path.splitext(os.path.basename(file_path))[
                    0] if file_path else "解析失败的文档",
                category='默认',
                file_type=file_ext,
                file_path=relative_path,
                status='草稿',
                saved_status='未存稿',
                likes=0,
                uploader_phone=current_user.phone if current_user else None,
                word_count=0,
                author_id=current_user.id if current_user else None,
                content_html=f"<p>文档解析失败: {str(e)}</p>",
                images_info=None
            )

    def _extract_metadata(self, doc: Document, file_path: str) -> Dict[str, Any]:
        """提取文档元数据"""
        try:
            # 获取文档属性
            core_props = doc.core_properties

            # 尝试从文档属性获取标题
            title = core_props.title or ""

            # 如果没有标题，使用第一段作为标题
            if not title and doc.paragraphs:
                first_paragraph = doc.paragraphs[0].text.strip()
                if first_paragraph:
                    title = first_paragraph[:50] + \
                            ("..." if len(first_paragraph) > 50 else "")

            # 如果还是没有标题，使用文件名
            if not title:
                title = os.path.splitext(os.path.basename(file_path))[0]

            metadata = {
                'title': title,
                'author': core_props.author or "",
                'subject': core_props.subject or "",
                'keywords': core_props.keywords or "",
                'comments': core_props.comments or "",
                'category': core_props.category or "默认",
                'created': core_props.created or datetime.now(),
                'modified': core_props.modified or datetime.now(),
                'last_modified_by': core_props.last_modified_by or "",
                'revision': core_props.revision or 1,
                'version': core_props.version or "",
                'file_name': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path)
            }

            return metadata

        except Exception as e:
            print(f"提取元数据失败: {str(e)}")
            return {
                'title': os.path.splitext(os.path.basename(file_path))[0],
                'author': "",
                'subject': "",
                'keywords': "",
                'comments': "",
                'category': "默认",
                'created': datetime.now(),
                'modified': datetime.now(),
                'file_name': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
            }

    def _extract_content_and_images(self, doc: Document, file_path: str) -> Tuple[str, List[Dict]]:
        """提取内容和图片，保持顺序"""
        try:
            html_parts = []
            images_info = []

            # 提取文档中的图片关系
            image_relations = self._get_image_relations(file_path)
            print(f"找到 {len(image_relations)} 张图片在文档中")

            # 遍历所有段落和元素
            for element in doc.element.body:
                # 处理段落
                if element.tag.endswith('}p'):  # w:p 段落元素
                    try:
                        # 从element创建段落对象
                        para = None
                        for p in doc.paragraphs:
                            if p._element == element:
                                para = p
                                break

                        if para:
                            # 先检查段落中是否有图片
                            para_images = self._extract_paragraph_images(
                                para, image_relations, file_path)

                            # 处理段落文本
                            para_html = self._process_paragraph(para)

                            # 如果段落有图片，需要将图片插入到正确位置
                            if para_images:
                                # 如果段落有文本内容，先添加文本
                                if para_html and para.text.strip():
                                    html_parts.append(para_html)

                                # 然后添加图片
                                for img_info in para_images:
                                    images_info.append(img_info)
                                    img_html = f'<img src="/{img_info["local_path"]}" alt="{img_info["alt_text"]}" style="max-width: 100%; height: auto; margin: 10px 0;">'
                                    html_parts.append(img_html)
                            elif para_html:
                                # 没有图片的普通段落
                                html_parts.append(para_html)

                    except Exception as e:
                        print(f"处理段落失败: {str(e)}")
                        continue

                # 处理表格
                elif element.tag.endswith('}tbl'):  # w:tbl 表格元素
                    try:
                        # 从element找到对应的表格对象
                        for table in doc.tables:
                            if table._element == element:
                                table_html = self._process_table(table)
                                if table_html:
                                    html_parts.append(table_html)
                                break
                    except Exception as e:
                        print(f"处理表格失败: {str(e)}")
                        continue

            # 确保所有在image_relations中的图片都被处理
            # 检查是否有未处理的图片（通过原始文件名）
            processed_images = set()
            for img in images_info:
                if 'original_filename' in img:
                    processed_images.add(img['original_filename'])
                else:
                    # 兼容旧版本，从filename推断
                    processed_images.add(img['filename'])

            unprocessed_images = set(image_relations.keys()) - processed_images

            if unprocessed_images:
                print(f"发现 {len(unprocessed_images)} 张未处理的图片，正在处理...")
                for img_name in unprocessed_images:
                    try:
                        img_data = image_relations[img_name]
                        ext = os.path.splitext(img_name.lower())[1]
                        content_type_map = {
                            '.png': 'image/png',
                            '.jpg': 'image/jpeg',
                            '.jpeg': 'image/jpeg',
                            '.gif': 'image/gif',
                            '.bmp': 'image/bmp',
                            '.tiff': 'image/tiff'
                        }
                        content_type = content_type_map.get(ext, 'image/png')

                        image_info = self._save_image_to_local(
                            img_data,
                            content_type,
                            alt_text=f"图片_{len(images_info) + 1}",
                            original_name=img_name
                        )
                        if image_info:
                            images_info.append(image_info)
                            # 在内容末尾添加图片
                            img_html = f'<img src="/{image_info["local_path"]}" alt="{image_info["alt_text"]}" style="max-width: 100%; height: auto; margin: 10px 0;">'
                            html_parts.append(img_html)
                            print(
                                f"成功处理图片: {img_name} -> {image_info['local_path']}")
                    except Exception as e:
                        print(f"处理未处理图片失败 {img_name}: {str(e)}")
                        continue

            content_html = '\n'.join(html_parts)
            print(f"提取完成: {len(images_info)} 张图片, HTML长度: {len(content_html)}")

            return content_html, images_info

        except Exception as e:
            print(f"提取内容失败: {str(e)}")
            return "", []

    def _get_image_relations(self, file_path: str) -> Dict[str, bytes]:
        """获取文档中的图片关系映射"""
        try:
            image_relations = {}

            # 读取docx文件（实际上是zip文件）
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # 查找所有可能的媒体文件路径
                media_paths = [
                    'word/media/',  # 标准路径
                    'word/embeddings/',  # 嵌入对象路径
                    'customXml/',  # 自定义XML中的媒体
                ]

                for file_info in docx_zip.filelist:
                    filename = file_info.filename

                    # 检查是否是媒体文件
                    is_media = any(filename.startswith(path)
                                   for path in media_paths)

                    if is_media:
                        # 检查文件扩展名是否是图片
                        ext = os.path.splitext(filename.lower())[1]
                        if ext in self.image_extensions:
                            try:
                                # 读取图片数据
                                image_data = docx_zip.read(filename)
                                image_name = os.path.basename(filename)
                                image_relations[image_name] = image_data
                                print(
                                    f"找到图片: {image_name} ({len(image_data)} bytes)")
                            except Exception as e:
                                print(f"读取图片文件失败 {filename}: {str(e)}")

            return image_relations

        except Exception as e:
            print(f"获取图片关系失败: {str(e)}")
            return {}

    def _extract_paragraph_images(self, paragraph, image_relations: Dict[str, bytes], file_path: str) -> List[Dict]:
        """提取段落中的图片"""
        images = []

        try:
            # 方法1: 通过drawing元素和relationship提取图片
            for run in paragraph.runs:
                # 查找图片元素
                for drawing in run._element.findall('.//a:blip',
                                                    {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                    embed = drawing.get(qn('r:embed'))
                    if embed:
                        try:
                            rel = paragraph.part.rels[embed]
                            image_part = rel.target_part

                            if hasattr(image_part, 'blob'):
                                # 保存图片到本地
                                image_info = self._save_image_to_local(
                                    image_part.blob,
                                    image_part.content_type,
                                    alt_text=f"图片_{len(images) + 1}",
                                    original_name=getattr(image_part, 'partname', '').split('/')[-1] if hasattr(
                                        image_part, 'partname') else ""
                                )
                                if image_info:
                                    images.append(image_info)
                        except Exception as e:
                            print(f"处理图片失败: {str(e)}")
                            continue

                # 方法2: 查找w:drawing元素（Word 2016+格式）
                for drawing in run._element.findall('.//w:drawing', {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    for inline in drawing.findall('.//wp:inline', {
                        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}):
                        for blip in inline.findall('.//a:blip',
                                                   {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                            embed = blip.get(qn('r:embed'))
                            if embed:
                                try:
                                    rel = paragraph.part.rels[embed]
                                    image_part = rel.target_part

                                    if hasattr(image_part, 'blob'):
                                        image_info = self._save_image_to_local(
                                            image_part.blob,
                                            image_part.content_type,
                                            alt_text=f"图片_{len(images) + 1}",
                                            original_name=getattr(image_part, 'partname', '').split('/')[-1] if hasattr(
                                                image_part, 'partname') else ""
                                        )
                                        if image_info:
                                            images.append(image_info)
                                except Exception as e:
                                    print(f"处理inline图片失败: {str(e)}")
                                    continue

            # 方法3: 如果前面的方法没有找到图片，尝试从image_relations中按顺序处理
            if not images and image_relations:
                for img_name, img_data in image_relations.items():
                    try:
                        # 根据文件扩展名推断content_type
                        ext = os.path.splitext(img_name.lower())[1]
                        content_type_map = {
                            '.png': 'image/png',
                            '.jpg': 'image/jpeg',
                            '.jpeg': 'image/jpeg',
                            '.gif': 'image/gif',
                            '.bmp': 'image/bmp',
                            '.tiff': 'image/tiff'
                        }
                        content_type = content_type_map.get(ext, 'image/png')

                        image_info = self._save_image_to_local(
                            img_data,
                            content_type,
                            alt_text=f"图片_{len(images) + 1}"
                        )
                        if image_info:
                            images.append(image_info)
                    except Exception as e:
                        print(f"处理image_relations图片失败: {str(e)}")
                        continue

        except Exception as e:
            print(f"提取段落图片失败: {str(e)}")

        return images

    def _save_image_to_local(self, image_data: bytes, content_type: str,
                            alt_text: str = "", original_name: str = "") -> Optional[Dict]:
        """保存图片到本地"""
        try:

            # 根据content_type确定文件扩展名
            extension_map = {
                'image/png': '.png',
                'image/jpeg': '.jpg',
                'image/jpg': '.jpg',
                'image/gif': '.gif',
                'image/bmp': '.bmp',
                'image/tiff': '.tiff'
            }

            extension = extension_map.get(content_type, '.png')

            # 生成唯一文件名
            unique_id = str(uuid.uuid4())
            filename = f"{unique_id}{extension}"
            local_path = os.path.join(self.images_dir, filename)

            # 保存图片
            with open(local_path, 'wb') as f:
                f.write(image_data)

            # 获取图片信息
            try:
                with Image.open(BytesIO(image_data)) as img:
                    width, height = img.size
                    format_name = img.format
            except Exception:
                width, height = 0, 0
                format_name = "unknown"

            # 从self.images_dir提取相对路径部分
            relative_path = f"{self.images_dir}/{filename}"

            image_info = {
                'filename': filename,
                'local_path': relative_path,  # 相对路径
                'absolute_path': local_path,  # 绝对路径
                'alt_text': alt_text,
                'content_type': content_type,
                'size': len(image_data),
                'width': width,
                'height': height,
                'format': format_name,
                'original_filename': original_name or filename  # 原始文件名（用于去重检查）
            }

            print(f"图片保存成功: {filename} (原始: {original_name})")
            return image_info

        except Exception as e:
            print(f"保存图片失败: {str(e)}")
            return None

    def _process_paragraph(self, paragraph) -> str:
        """处理段落，转换为HTML"""
        try:
            if not paragraph.text.strip():
                return ""

            # 获取段落样式
            style_name = paragraph.style.name if paragraph.style else "Normal"

            # 处理标题
            if style_name.startswith('Heading'):
                level = style_name.replace('Heading ', '') or '1'
                try:
                    level = int(level)
                    level = min(max(level, 1), 6)  # 限制在1-6之间
                except:
                    level = 1

                content = self._process_runs(paragraph.runs)
                return f"<h{level}>{content}</h{level}>"

            # 处理普通段落
            content = self._process_runs(paragraph.runs)
            if not content:
                return ""

            # 处理段落对齐
            alignment = ""
            if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                alignment = ' style="text-align: center;"'
            elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                alignment = ' style="text-align: right;"'
            elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                alignment = ' style="text-align: justify;"'

            return f"<p{alignment}>{content}</p>"

        except Exception as e:
            print(f"处理段落失败: {str(e)}")
            return ""

    def _process_runs(self, runs) -> str:
        """处理文本运行，应用格式"""
        try:
            result = ""

            for run in runs:
                text = run.text
                if not text:
                    continue

                # 转义HTML特殊字符
                text = text.replace('&', '&amp;')
                text = text.replace('<', '&lt;')
                text = text.replace('>', '&gt;')
                text = text.replace('"', '&quot;')
                text = text.replace("'", '&#39;')

                # 应用格式
                if run.bold:
                    text = f"<strong>{text}</strong>"
                if run.italic:
                    text = f"<em>{text}</em>"
                if run.underline:
                    text = f"<u>{text}</u>"

                # 处理字体颜色
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    color = f"#{rgb.red:02x}{rgb.green:02x}{rgb.blue:02x}"
                    text = f'<span style="color: {color};">{text}</span>'

                # 处理字体大小
                if run.font.size:
                    size_pt = run.font.size.pt
                    text = f'<span style="font-size: {size_pt}pt;">{text}</span>'

                result += text

            return result

        except Exception as e:
            print(f"处理文本运行失败: {str(e)}")
            return ""

    def _process_table(self, table) -> str:
        """处理表格，转换为HTML"""
        try:
            html = ['<table class="table table-bordered">']

            for i, row in enumerate(table.rows):
                html.append('<tr>')

                for cell in row.cells:
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        para_content = self._process_runs(paragraph.runs)
                        if para_content:
                            cell_text += para_content + "<br>"

                    # 移除最后的<br>
                    cell_text = cell_text.rstrip("<br>")

                    # 如果是第一行，使用th标签
                    tag = 'th' if i == 0 else 'td'
                    html.append(f'<{tag}>{cell_text}</{tag}>')

                html.append('</tr>')

            html.append('</table>')
            return '\n'.join(html)

        except Exception as e:
            print(f"处理表格失败: {str(e)}")
            return ""

    def _calculate_word_count(self, html_content: str) -> int:
        """计算字数（去除HTML标签）"""
        try:
            # 简单的HTML标签移除
            text = re.sub(r'<[^>]+>', '', html_content)
            # 计算中文字符和英文单词
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
            english_words = len(re.findall(r'\b[a-zA-Z]+\b', text))
            return chinese_chars + english_words
        except:
            return 0


# 创建全局实例
docx_process = DocxProcess()
